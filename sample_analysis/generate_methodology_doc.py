"""
generate_methodology_doc.py

Generate the Word (.docx) methodology document for the R&D / science team,
describing how dose analysis is performed on data returned from the ISS (in the
LeadEnclosureData.csv format).

The calibration table and shielding map are pulled directly from calibration.py
so the document can never disagree with the code.

Run:
    cd sample_analysis
    python generate_methodology_doc.py
Writes: ../RADFET_ISS_Dose_Analysis_Methodology.docx
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from calibration import (
    CALIBRATION_CURVES,
    SHIELDING_BY_CHANNEL,
    TRIAL_GROUPS,
    UNSHIELDED_CHANNEL,
)

SCRIPT_DIR = Path(__file__).resolve().parent
OUTPUT_DOCX = SCRIPT_DIR.parent / "RADFET_ISS_Dose_Analysis_Methodology.docx"

ACCENT = RGBColor(0x1F, 0x4E, 0x79)


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    return h


def add_body(doc, text):
    p = doc.add_paragraph(text)
    return p


def add_bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")


def add_numbered(doc, text):
    return doc.add_paragraph(text, style="List Number")


def add_mono(doc, text):
    """A monospace block for formulae / pseudocode."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    return p


def build() -> None:
    doc = Document()

    # Base style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # ---- Title block -----------------------------------------------------
    title = doc.add_heading("RADFET ISS Dose-Analysis Methodology", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    sub = doc.add_paragraph()
    r = sub.add_run("Converting flight dosimeter telemetry to absorbed dose, "
                    "with shielding comparison and uncertainty")
    r.italic = True
    r.font.color.rgb = ACCENT

    meta = doc.add_paragraph()
    meta.add_run("Audience: ").bold = True
    meta.add_run("R&D / Science team\n")
    meta.add_run("Prepared: ").bold = True
    meta.add_run("2026-06-17\n")
    meta.add_run("Status: ").bold = True
    meta.add_run("Draft for review")

    doc.add_paragraph(
        "This document describes how we will analyze the dosimetry data "
        "returned from the ISS. The flight data arrives in the same CSV format "
        "as LeadEnclosureData.csv. A runnable reference implementation of every "
        "step below lives in the sample_analysis/ folder and is demonstrated on "
        "synthetic data."
    )

    # ---- 1. Purpose & scope ---------------------------------------------
    add_heading(doc, "1. Purpose and scope", 1)
    add_body(doc,
             "We fly five RADFET dosimeters, each behind a different shielding "
             "stack, to measure how effectively each stack attenuates the "
             "on-orbit radiation dose. This document defines the agreed "
             "procedure for turning the raw voltage telemetry into absorbed "
             "dose (in Rad), for combining repeat measurements, for quantifying "
             "uncertainty, and for comparing the shielding configurations.")

    # ---- 2. Input data format -------------------------------------------
    add_heading(doc, "2. Input data format", 1)
    add_body(doc,
             "Each row is one channel read at one time. Columns (identical to "
             "LeadEnclosureData.csv):")
    cols = [
        ("timestamp", "ISO-8601 read time (host clock)."),
        ("sensor_group", "R1 or R2 - the two read-out groups (see Section 4)."),
        ("channel", "1-5 - identifies the sensor / shielding configuration."),
        ("raw_adc", "12-bit ADC count (0-4095) behind the voltage."),
        ("voltage", "Threshold-voltage shift dVt in Volts (already computed)."),
        ("dose_rad", "Convenience dose column; the analysis recomputes dose itself."),
        ("raw_timestamp", "Secondary timestamp; may contain epoch corruption."),
    ]
    t = doc.add_table(rows=1, cols=2)
    t.style = "Light Grid Accent 1"
    t.rows[0].cells[0].paragraphs[0].add_run("Column").bold = True
    t.rows[0].cells[1].paragraphs[0].add_run("Meaning").bold = True
    for name, meaning in cols:
        cells = t.add_row().cells
        cells[0].paragraphs[0].add_run(name).font.name = "Consolas"
        cells[1].text = meaning
    add_body(doc,
             "Important: the voltage column already carries the threshold-voltage "
             "shift dVt (irradiated minus non-irradiated reference), so no extra "
             "baseline subtraction is performed in the dose conversion.")

    # ---- 3. Calibration model -------------------------------------------
    add_heading(doc, "3. Calibration model", 1)
    add_body(doc,
             "We use the Varadis QF 16 RADFET calibration (400nm IMPL RADFET, "
             "plastic package; Mask-set COMRAD; Lot P5925-W3; Co-60 source at "
             "5.0 kRad/hour; Issue No.01, 01/08/2023). The fit relates the "
             "threshold-voltage shift to absorbed dose by a power law:")
    add_mono(doc, "    dVt [V] = A * Dose[Rad] ^ B")
    add_body(doc, "Inverting to recover dose from a measured dVt:")
    add_mono(doc, "    Dose[Rad] = ( dVt / A ) ^ (1 / B)")
    add_body(doc,
             "Varadis provides FIVE fits of the same data, each valid over a "
             "different dose range, and advises using the curve applicable to "
             "the actual dose. The narrow fits are the most accurate at low "
             "dose. Coefficients (single source of truth: calibration.py):")

    t = doc.add_table(rows=1, cols=6)
    t.style = "Light Grid Accent 1"
    for i, head in enumerate(["Dose range", "A", "sigma(A)", "B", "sigma(B)", "R^2"]):
        t.rows[0].cells[i].paragraphs[0].add_run(head).bold = True
    for c in CALIBRATION_CURVES:
        cells = t.add_row().cells
        cells[0].text = c.name
        cells[1].text = f"{c.A:.4g}"
        cells[2].text = f"{c.sigma_A:.3e}"
        cells[3].text = f"{c.B:.4g}"
        cells[4].text = f"{c.sigma_B:.3e}"
        cells[5].text = f"{c.r_square:.3f}"

    # ---- 4. Sensors & shielding -----------------------------------------
    add_heading(doc, "4. Sensors, shielding, and trials", 1)
    add_body(doc,
             "Five sensors are flown, one per channel, each behind a different "
             "shielding stack:")
    t = doc.add_table(rows=1, cols=2)
    t.style = "Light Grid Accent 1"
    t.rows[0].cells[0].paragraphs[0].add_run("Channel").bold = True
    t.rows[0].cells[1].paragraphs[0].add_run("Shielding configuration").bold = True
    for ch in sorted(SHIELDING_BY_CHANNEL):
        cells = t.add_row().cells
        cells[0].text = str(ch)
        label = SHIELDING_BY_CHANNEL[ch]
        if ch == UNSHIELDED_CHANNEL:
            label += "  (reference for attenuation)"
        cells[1].text = label
    add_body(doc,
             f"Each sensor is read out at two positions/units recorded as "
             f"sensor_group {' and '.join(TRIAL_GROUPS)}. We treat R1 and R2 as "
             f"two independent TRIALS (replicates) of the same shielding "
             f"configuration. This gives, per configuration, two dose estimates "
             f"whose agreement is itself a check on repeatability.")

    # ---- 5. Range selection ---------------------------------------------
    add_heading(doc, "5. Calibration-curve (range) selection", 1)
    add_body(doc,
             "Because the correct curve depends on the dose we are trying to "
             "measure, we select it self-consistently. The five ranges are "
             "nested (0-1 kRad within 0-5 kRad within ... within 0-100 kRad), "
             "so we walk from the narrowest curve to the widest and accept the "
             "first curve whose own dose estimate lands inside its validity "
             "range:")
    add_mono(doc,
             "for curve in [0-1, 0-5, 0-10, 0-50, 0-100] kRad:   # narrow -> wide\n"
             "    dose = (dVt / curve.A) ^ (1 / curve.B)\n"
             "    if dose <= curve.dose_max:\n"
             "        return curve, dose          # most accurate valid fit\n"
             "# if none qualify, report the 0-100 kRad estimate, flagged EXTRAP")
    add_body(doc,
             "A reading whose dose exceeds even the 0-100 kRad fit is flagged as "
             "an extrapolation rather than silently trusted.")

    # ---- 6. Data quality -------------------------------------------------
    add_heading(doc, "6. Data-quality control", 1)
    add_body(doc,
             "The real telemetry contains corrupted rows. We never silently drop "
             "them; we flag them and exclude them from statistics so the dose "
             "reflects genuine sensor behavior. A reading is excluded when:")
    add_bullet(doc, "the voltage is non-finite or outside the plausible band "
                    "(e.g. the ~3.8e14 corruption seen in real data);")
    add_bullet(doc, "the ADC is at/near full scale (saturation / rail hit), such "
                    "as the synchronous spike where all channels read ~max for "
                    "one cycle;")
    add_bullet(doc, "the timestamp is out of range (e.g. 1969-epoch rows). Such "
                    "timestamps are treated as missing so they cannot masquerade "
                    "as the first reading; the voltage may still be usable.")
    add_body(doc,
             "Counts of each rejection category are reported so data quality is "
             "visible, not hidden.")

    # ---- 7. Aggregation & combination -----------------------------------
    add_heading(doc, "7. Per-sensor aggregation and trial combination", 1)
    add_numbered(doc, "For each (sensor_group, channel), average dVt over the "
                      "valid readings to get the sensor's mean dVt. The "
                      "measurement uncertainty on that mean is the characterized "
                      "per-sensor sigma_V from the lead-brick run (Section 8), "
                      "not an SEM re-estimated from flight data.")
    add_numbered(doc, "Combine the R1 and R2 trials per channel by "
                      "inverse-variance weighting in voltage space, giving a "
                      "combined dVt and its combined sigma_V. (Half the spread "
                      "between the two trial doses is also reported as a "
                      "repeatability check.)")
    add_numbered(doc, "Convert the combined dVt to dose using the selected curve "
                      "(Section 5), and attach the full propagated uncertainty "
                      "(Section 8).")

    # ---- 8. Uncertainty --------------------------------------------------
    add_heading(doc, "8. Uncertainty quantification", 1)
    add_body(doc,
             "Per-sensor measurement uncertainty (sigma_V). In the lead-brick "
             "run the dose is essentially fixed, so each sensor's voltage spread "
             "is its characterized measurement noise floor. We take the "
             "per-sensor 'std dev (sample)' values directly from "
             "sensor_analysis/analysis_report.txt (read at runtime, so they stay "
             "in sync if the lead-brick analysis is re-run). These are of order "
             "0.004-0.006 V and differ slightly per sensor. A coverage factor "
             "(e.g. 2-sigma) can be applied centrally if a more conservative "
             "budget is wanted.")
    add_body(doc,
             "Absolute dose uncertainty. We propagate a 1-sigma dose uncertainty "
             "through the inversion Dose = (dVt/A)^(1/B), combining the measured "
             "sigma_V with the calibration uncertainties sigma(A) and sigma(B):")
    add_mono(doc,
             "(sigma_Dose / Dose)^2 =\n"
             "      ( sigma_V    / (B * dVt) )^2    # measurement (lead-brick sigma_V)\n"
             "    + ( sigma_A   / (B * A)   )^2    # calibration uncertainty in A\n"
             "    + ( ln(Dose) * sigma_B / B )^2   # calibration uncertainty in B")
    add_body(doc,
             "This full sigma is the error bar on the ABSOLUTE dose of a single "
             "configuration. The two trials R1 and R2 are combined by "
             "inverse-variance weighting in voltage space; the spread between "
             "them is reported separately as a repeatability check.")

    # ---- 9. Shielding effectiveness & significance ----------------------
    add_heading(doc, "9. Shielding effectiveness and statistical significance", 1)
    add_body(doc,
             "Using the bare sensor (channel "
             f"{UNSHIELDED_CHANNEL}) as the reference, each shielded "
             "configuration is summarized by:")
    add_bullet(doc, "Attenuation factor = dose(bare) / dose(config)  "
                    "(higher means more effective shielding).")
    add_bullet(doc, "Dose reduction (%) = (1 - dose(config) / dose(bare)) x 100.")

    add_body(doc,
             "Is the difference statistically significant? This is the key "
             "question and it has an important subtlety. The same calibration "
             "curve is applied to every sensor, so the calibration "
             "uncertainties sigma(A), sigma(B) are COMMON-MODE: they shift all "
             "configurations together and largely cancel in a DIFFERENCE between "
             "two configurations. Adding them into a significance test would "
             "make us under-confident. We therefore test significance in "
             "VOLTAGE space, using only the measured sigma_V:")
    add_mono(doc,
             "z = ( dVt_bare - dVt_config ) / sqrt( sigma_V_bare^2 + sigma_V_config^2 )")
    add_body(doc,
             "Because dose is a monotonic function of dVt, a significant "
             "difference in dVt is a significant difference in dose, and this "
             "avoids entangling the (correlated) calibration error. The combined "
             "sigma_V per configuration comes from the inverse-variance "
             "combination of its R1 and R2 trials. We flag |z| > 1.96 as "
             "significant at 95% (~2 sigma) and |z| > 3 at ~99.7% (~3 sigma), "
             "and also report the two-sided p-value.")
    add_body(doc,
             "Outputs include each configuration tested against the bare sensor "
             "and a full pairwise |z| matrix across all five configurations, so "
             "neighbouring stacks (which may differ by little) can be compared "
             "directly. The absolute dose with its full uncertainty (Section 8) "
             "is reported alongside for magnitude.")

    # ---- 10. Outputs -----------------------------------------------------
    add_heading(doc, "10. Outputs", 1)
    add_bullet(doc, "A text report: data-quality summary, per-sensor / per-trial "
                    "dose, and the shielding-effectiveness table.")
    add_bullet(doc, "A tidy CSV: one row per shielding configuration (dose, "
                    "uncertainties, attenuation factor, % reduction) for "
                    "downstream plotting and reporting.")

    # ---- 11. Reference implementation -----------------------------------
    add_heading(doc, "11. Reference implementation", 1)
    add_body(doc, "The sample_analysis/ folder contains a runnable demonstration "
                  "on synthetic data:")
    add_bullet(doc, "calibration.py - calibration coefficients, range selection, "
                    "uncertainty, and the channel->shielding map (single source "
                    "of truth; this document is generated from it).")
    add_bullet(doc, "generate_fake_data.py - builds an ISS-format dataset with "
                    "known true doses and injected anomalies.")
    add_bullet(doc, "analyze_sample.py - runs the full pipeline; on real data "
                    "only the input path changes.")

    # ---- 12. Assumptions & open questions -------------------------------
    add_heading(doc, "12. Assumptions and open questions for review", 1)
    add_bullet(doc, "The voltage column is the threshold-voltage shift dVt "
                    "(already baseline-subtracted). Please confirm.")
    add_bullet(doc, "Channel->shielding mapping (Section 4) is as listed. "
                    "Confirm before the flight run.")
    add_bullet(doc, "The Co-60 ground calibration is assumed applicable to the "
                    "mixed ISS radiation environment; any spectral correction "
                    "factor is out of scope here and would be applied downstream.")
    add_bullet(doc, "Valid-voltage band and saturation threshold are tuned to "
                    "the lead-enclosure baseline; revisit for the higher dVt "
                    "expected on orbit.")

    doc.save(OUTPUT_DOCX)
    print(f"Wrote {OUTPUT_DOCX}")


if __name__ == "__main__":
    build()
